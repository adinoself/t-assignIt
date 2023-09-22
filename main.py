import requests
import streamlit as st
import transcribe
import time
import sys
from zipfile import ZipFile
from time import sleep
import os
import pickle
from pathlib import Path
import streamlit_authenticator as stauth
from docx import Document
from zipfile import ZipFile
import base64

auth_key = st.secrets['auth_key']

st.header("Songa Prime Transcription")
hide_st_style = """
            <style>
            #MainMenu {visibility: hidden;}
            footer {visibility: hidden;}
            header {visibility: hidden;}
            </style>
            """
st.markdown(hide_st_style, unsafe_allow_html=True)

# --- USER AUTHENTICATION ---
names = ["Joseph Modi", "Rael Orgut"]
usernames = ["adinoself", "raelorgut"]
credentials = {"usernames": {}}

# LOAD HASHED PASSWORDS
file_path = Path(__file__).parent / "hashed_pw.pkl"
with file_path.open("rb") as file:
    hashed_passwords = pickle.load(file)

authenticator = stauth.Authenticate(names, usernames, hashed_passwords, "transcription", "abcdef", cookie_expiry_days=30)

name, authentication_status, username = authenticator.login("Login", "main")

if authentication_status == False:
    st.error("Username/password is incorrect")

if authentication_status == None:
    st.warning("Please enter your username and password")

if authentication_status:

    # ----SIDEBAR ---
    authenticator.logout("Logout", "sidebar")
    st.sidebar.title(f"Welcome, {name}")

    # 1
    # THIS IS THE AUTHORIZATION HEADER
    def get_url(auth_key, data):
        '''
          Parameter:
            token: The API key
            data : The File Object to upload
          Return Value:
            url  : Url to uploaded file
        '''
        headers = {'authorization': auth_key}
        response = requests.post('https://api.assemblyai.com/v2/upload',
                                 headers=headers,
                                 data=data)
        url = response.json()["upload_url"]
        print("Uploaded File and got temporary URL to file")
        return url

    # 2 GET THE TRANSCRIBE ID
    def get_transcribe_id(auth_key, url):
        '''
          Parameter:
            token: The API key
            url  : Url to uploaded file
          Return Value:
            id   : The transcription id of the file
        '''
        endpoint = "https://api.assemblyai.com/v2/transcript"
        json = {
            "audio_url": url,
            "speaker_labels": True
        }
        headers = {
            "authorization": auth_key,
            "content-type": "application/json"
        }
        response = requests.post(endpoint, json=json, headers=headers)
        id = response.json()['id']
        print("Made request and file is currently queued")
        print(response.json())
        return id

    # 3 UPLOADED FILE
    def upload_file(fileObj):
        '''
          Parameter:
            fileObj: The File Object to transcribe
          Return Value:
            token  : The API key
            transcribe_id: The ID of the file which is being transcribed
        '''
        auth_key = st.secrets['auth_key']
        url = get_url(auth_key, fileObj)
        t_id = get_transcribe_id(auth_key, url)
        print("transcribe_id")
        return auth_key, t_id

    # 4 GET THE TRANSCRIPTION RESULT
    def get_text(auth_key, transcribe_id):
        '''
          Parameter:
            token: The API key
            transcribe_id: The ID of the file which is being
          Return Value:
            result : The response object
        '''
        endpoint = f"https://api.assemblyai.com/v2/transcript/{transcribe_id}"
        headers = {
            "authorization": auth_key,
        }
        result = requests.get(endpoint, headers=headers).json()
        #print(result.text)
        print(result)
        return result

    # 4.2 SPEAKER LABELS
    def get_speakers(speak):
        endpoint = f"https://api.assemblyai.com/v2/transcript/{transcribe_id}"
        headers = {
            "authorization": auth_key,
        }
        if speak['speaker'] == 'A':
            return "Speaker A: {text}".format(text=speak['text']) + "\n"

        if speak['speaker'] == 'B':
            return "Speaker B: {text}".format(text=speak['text']) + "\n"

        result = requests.get(endpoint, headers=headers).json()
        return result

    # ---- THE APP ---- FILE UPLOADER UI ----
    fileObject = st.file_uploader(label="Please upload your file")
    if fileObject:
        auth_key, transcribe_id = upload_file(fileObject)
        result = {}
        # polling
        sleep_duration = 1
        percent_complete = 0
        progress_bar = st.progress(percent_complete)
        st.text("Currently in queue")
        while result.get("status") != "processing":
            percent_complete += sleep_duration
            time.sleep(sleep_duration)
            progress_bar.progress(percent_complete / 10)
            result = get_text(auth_key, transcribe_id)

        sleep_duration = 0.01

        # CHECK IF TRANSCRIPTION IS COMPLETE

        for percent in range(percent_complete, 101):
            time.sleep(sleep_duration)
            progress_bar.progress(percent)

        with st.spinner("Processing....."):
            while result.get("status") != 'completed':
                # while result.json()['status'] != 'completed':
                # sleep(5)
                # st.warning('Transcription is processing...')
                result = get_text(auth_key, transcribe_id)
                # result = requests.get(endpoint, headers=headers)

        # PRINT TRANSCRIBED TEXT
        st.success('Transcription Successful!')
        st.subheader("Transcribed Text")
        st.success(result['text'])

        # SAVE TRANSCRIBED TEXT TO A FILE

        # SAVE AS A TXT FILE
        doc = Document()
        doc.add_paragraph(result["text"])
        doc.save('plain text transcript.docx')

        # SAVE AS A TXT FILE WITH SPEAKER LABELS
        doc = Document()
        for utterance in result["utterances"]:
            speaker = utterance["speaker"]
            text = utterance["text"]
            doc.add_paragraph(f"Speaker {speaker}: {text}")
        doc.save('transcript with speaker labels.docx')

        # SAVE AS SRT FILE
        endpoint = f"https://api.assemblyai.com/v2/transcript/{transcribe_id}"
        headers = {
            "authorization": auth_key,
        }
        result = requests.get(endpoint, headers=headers).json()
        srt_endpoint = endpoint + "/srt"
        srt_response = requests.get(srt_endpoint, headers=headers)
        with open("transcript with time stamps.txt", "w") as _file:
            _file.write(srt_response.text)
            print(srt_response)

        # Create a zip file
        with ZipFile('transcription.zip', 'w') as zip_file:
            zip_file.write('plain text transcript.docx')
            zip_file.write('transcript with speaker labels.docx')
            zip_file.write('transcript with time stamps.txt')

        # Generate base64 encoded data for the zip file
        with open("transcription.zip", "rb") as zip_file:
            zip_data = zip_file.read()
            zip_data_base64 = base64.b64encode(zip_data).decode()

        # Display download buttons
        st.subheader("Download Transcripts")
        col1, col2, col3 = st.columns(3)

        with col1:
            download_filename1 = "plain_text_transcript.docx"
            download_link1 = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{base64.b64encode(open("plain text transcript.docx", "rb").read()).decode()}" download="{download_filename1}">Download Plain Text Transcript.docx</a>'
            st.markdown(download_link1, unsafe_allow_html=True)

        with col2:
            download_filename2 = "transcript_with_speaker_labels.docx"
            download_link2 = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{base64.b64encode(open("transcript with speaker labels.docx", "rb").read()).decode()}" download="{download_filename2}">Download Transcript with Speaker Labels.docx</a>'
            st.markdown(download_link2, unsafe_allow_html=True)

        with col3:
            download_filename3 = "transcript_with_time_stamps.txt"
            download_link3 = f'<a href="data:text/plain;base64,{base64.b64encode(open("transcript with time stamps.txt", "rb").read()).decode()}" download="{download_filename3}">Download Transcript with Time Stamps.txt</a>'
            st.markdown(download_link3, unsafe_allow_html=True)

        # Display download buttons
        st.subheader("Download All Transcripts")

        # Display download zip button
        download_filename_zip = "transcription.zip"
        download_link_zip = f'<a href="data:application/zip;base64,{zip_data_base64}" download="{download_filename_zip}">Download All Transcripts as ZIP</a>'
        st.markdown(download_link_zip, unsafe_allow_html=True)